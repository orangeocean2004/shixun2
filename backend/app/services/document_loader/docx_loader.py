from __future__ import annotations

import os
import shutil
import zipfile
from pathlib import Path

from backend.app.services.segmenting.parser import detect_block_type
from backend.app.services.segmenting.models import DocumentBlock

from .exceptions import DocumentLoaderError

# 图片存储目录
_IMAGES_DIR = Path(__file__).resolve().parents[4] / "data" / "images"


def load_docx_file(file_path: str | Path, doc_id: str | None = None) -> list[DocumentBlock]:
    """读取 DOCX 文件，含段落、表格和图片提取。"""

    try:
        from docx import Document
    except ImportError as exc:
        raise DocumentLoaderError("读取 DOCX 需要安装 python-docx：pip install python-docx") from exc

    path = Path(file_path)
    document = Document(str(path))
    blocks: list[DocumentBlock] = []

    # 提取文档内嵌图片
    img_dir = doc_id or "".join(c if c.isalnum() else "_" for c in path.stem).strip("_") or "doc"
    image_map = _extract_docx_images(path, img_dir)
    image_index = 0

    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()

        # 检查段落是否包含图片
        has_image = _paragraph_has_image(paragraph)
        if has_image and image_index < len(image_map):
            img_filename = image_map[image_index]
            image_index += 1
            image_block_text = f"[IMAGE: {img_filename}]"
            if text:
                image_block_text = f"{text}\n{image_block_text}"
            blocks.append(
                DocumentBlock(
                    block_id=f"p{index:04d}",
                    text=image_block_text,
                    block_type="image",
                    metadata={"style": paragraph.style.name, "image": img_filename},
                )
            )
            continue

        if not text:
            continue

        block_type = _detect_docx_paragraph_type(paragraph.style.name, text)
        blocks.append(
            DocumentBlock(
                block_id=f"p{index:04d}",
                text=text,
                block_type=block_type,
                metadata={"style": paragraph.style.name},
            )
        )

    table_offset = len(blocks)
    for table_index, table in enumerate(document.tables):
        table_text = _table_to_markdown(table)
        if not table_text:
            continue
        blocks.append(
            DocumentBlock(
                block_id=f"t{table_offset + table_index:04d}",
                text=table_text,
                block_type="table",
                metadata={"table_index": table_index},
            )
        )

    return blocks


def _extract_docx_images(docx_path: Path, doc_id: str) -> list[str]:
    """从 DOCX 中提取所有内嵌图片，返回文件名列表（按出现顺序）。"""
    filenames: list[str] = []
    try:
        doc_dir = _IMAGES_DIR / doc_id
        doc_dir.mkdir(parents=True, exist_ok=True)

        with zipfile.ZipFile(docx_path, "r") as zf:
            # DOCX 图片存储在 word/media/ 下
            image_entries = sorted(
                [name for name in zf.namelist() if name.startswith("word/media/")],
            )
            for i, entry_name in enumerate(image_entries):
                ext = os.path.splitext(entry_name)[1] or ".png"
                filename = f"img_{i:03d}{ext}"
                dest_path = doc_dir / filename
                if not dest_path.exists():
                    with zf.open(entry_name) as src:
                        with open(dest_path, "wb") as dst:
                            shutil.copyfileobj(src, dst)
                filenames.append(filename)
    except (zipfile.BadZipFile, OSError):
        pass
    return filenames


def _paragraph_has_image(paragraph) -> bool:
    """检查 DOCX 段落是否包含内嵌图片。"""
    try:
        # python-docx: 检查段落 XML 中是否包含图片引用
        nsmap = {"wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"}
        drawings = paragraph._element.findall(".//wp:inline", nsmap)
        if not drawings:
            drawings = paragraph._element.findall(".//wp:anchor", nsmap)
        return len(drawings) > 0
    except Exception:
        return False


def _detect_docx_paragraph_type(style_name: str, text: str) -> str:
    """根据 DOCX 样式和文本规则识别段落类型。"""

    if style_name.lower().startswith("heading"):
        return "heading"
    return detect_block_type(text)


def _table_to_markdown(table: object) -> str:
    """把 DOCX 表格临时转成 Markdown 表格，方便后续特殊块保护。"""

    rows: list[list[str]] = []
    for row in table.rows:
        values = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        if any(values):
            rows.append(values)

    if not rows:
        return ""

    width = max(len(row) for row in rows)
    normalized_rows = [row + [""] * (width - len(row)) for row in rows]
    header = normalized_rows[0]
    separator = ["---"] * width
    body = normalized_rows[1:]

    markdown_rows = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(separator) + " |",
    ]
    markdown_rows.extend("| " + " | ".join(row) + " |" for row in body)
    return "\n".join(markdown_rows)
